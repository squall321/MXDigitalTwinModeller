"""
MX Digital Twin Modeller - SpaceClaim 기능 레포트 생성기
Word (.docx) 문서 생성 (한국어 폰트 지원)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ========================================================================
# 한국어 폰트 설정 헬퍼
# ========================================================================

def set_run_font(run, font_name='맑은 고딕', size=10, bold=False, color=None):
    """런에 한국어 폰트 적용"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 동아시아 폰트 설정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_heading_kr(doc, text, level=1):
    """한국어 호환 헤딩 추가"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_run_font(run, '맑은 고딕', size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True)
    return heading


def add_paragraph_kr(doc, text, bold=False, size=10, color=None, space_after=6):
    """한국어 호환 본문 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, '맑은 고딕', size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet_kr(doc, text, level=0, bold_prefix='', size=10):
    """한국어 호환 불릿 추가"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, '맑은 고딕', size=size, bold=True)
        run = p.add_run(text)
        set_run_font(run, '맑은 고딕', size=size)
    else:
        run = p.add_run(text)
        set_run_font(run, '맑은 고딕', size=size)
    return p


def add_table_kr(doc, headers, rows, col_widths=None):
    """한국어 호환 테이블 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, '맑은 고딕', size=9, bold=True, color=(255, 255, 255))
        # 헤더 배경색
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '2E5090'
        })
        shading.append(shading_elm)

    # 데이터 행
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, '맑은 고딕', size=9)

    # 컬럼 너비 설정
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def add_screenshot_placeholder(doc, caption=''):
    """스크린샷 자리표시자"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ 스크린샷: {caption} ]')
    set_run_font(run, '맑은 고딕', size=10, color=(150, 150, 150))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 박스 효과 (회색 배경)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0'
    })
    pPr.append(shd)
    return p


# ========================================================================
# 문서 생성
# ========================================================================

def create_report():
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ====================================================================
    # 표지
    # ====================================================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('MX Digital Twin Modeller')
    set_run_font(run, '맑은 고딕', size=28, bold=True, color=(46, 80, 144))

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('SpaceClaim Add-In 기능 레포트')
    set_run_font(run, '맑은 고딕', size=18, color=(80, 80, 80))

    doc.add_paragraph()

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_p.add_run('Version 1.1.0')
    set_run_font(run, '맑은 고딕', size=14, color=(120, 120, 120))

    doc.add_paragraph()

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_p.add_run('MX Composites')
    set_run_font(run, '맑은 고딕', size=14, bold=True, color=(46, 80, 144))

    from datetime import date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(date.today().strftime('%Y년 %m월 %d일'))
    set_run_font(run, '맑은 고딕', size=12, color=(120, 120, 120))

    doc.add_page_break()

    # ====================================================================
    # 목차
    # ====================================================================
    add_heading_kr(doc, '목차', level=1)
    toc_items = [
        '1. 개요',
        '2. 시스템 요구사항',
        '3. Ribbon UI 구성',
        '4. Specimen 그룹 (시편 생성)',
        '    4.1 인장시편 (Tensile Specimen)',
        '    4.2 굽힘시편 (Bending Specimen)',
        '    4.3 벤딩 지그 (Bending Fixture)',
        '    4.4 압축시편 (Compression Specimen)',
        '5. Advanced 그룹 (고급 시험)',
        '    5.1 CAI시편 (Compression After Impact)',
        '    5.2 피로시편 (Fatigue Specimen)',
        '    5.3 접합부시편 (Joint Specimen)',
        '6. Parametric 그룹 (파라메트릭 모델러)',
        '    6.1 적층 모델 (Laminate)',
        '    6.2 Void Cut (보이드 컷)',
        '7. Mesh 그룹 (메쉬 및 해석 설정)',
        '    7.1 Material (재료 물성)',
        '    7.2 Mesh 설정',
        '    7.3 Mesh Export (메쉬 내보내기)',
        '    7.4 Simplify (바디 단순화)',
        '    7.5 Contact (접촉면 감지)',
        '    7.6 STEP Export',
        '    7.7 Load (하중 정의)',
        '    7.8 Simulation Setup (시뮬레이션 설정)',
        '    7.9 Pipeline (일괄 실행)',
        '    7.10 Conformal Mesh (적합 메쉬)',
    ]
    for item in toc_items:
        add_paragraph_kr(doc, item, size=10, space_after=2)

    doc.add_page_break()

    # ====================================================================
    # 1. 개요
    # ====================================================================
    add_heading_kr(doc, '1. 개요', level=1)
    add_paragraph_kr(doc, (
        'MX Digital Twin Modeller는 ANSYS SpaceClaim용 Add-In으로, '
        '복합재 시편 모델링부터 메쉬 생성, 해석 설정, 내보내기까지의 '
        '전체 워크플로를 자동화하는 통합 도구입니다.'
    ))
    add_paragraph_kr(doc, (
        '본 레포트는 SpaceClaim Add-In에 포함된 모든 기능을 설명합니다. '
        '각 기능에 대해 지원 규격, 파라미터, 동작 방식을 상세히 기술합니다.'
    ))

    add_heading_kr(doc, '주요 특징', level=3)
    features = [
        '국제 규격 기반 시편 자동 생성 (ASTM, ISO, DMA, Boeing 등 40+ 규격)',
        '파라메트릭 적층 복합재 모델링 (Rectangular, Surface, Solid 모드)',
        'Boolean 연산 기반 Void Cut 기능',
        '자동 메쉬 생성 및 바디별 방향성 크기 제어',
        '접촉면 자동 감지 및 Named Selection 생성',
        'LS-DYNA, ANSYS, Abaqus, Fluent, CGNS 형식 메쉬 내보내기',
        '하중 정의 (수식, 테이블, PWM 모드) 및 시뮬레이션 설정',
        '일괄 실행 파이프라인 (단순화→재료→접촉→메쉬→내보내기)',
        'STEP 어셈블리 기반 Conformal Mesh 생성',
        'Modeless 대화창 + 실시간 미리보기'
    ]
    for f in features:
        add_bullet_kr(doc, f)

    doc.add_page_break()

    # ====================================================================
    # 2. 시스템 요구사항
    # ====================================================================
    add_heading_kr(doc, '2. 시스템 요구사항', level=1)
    add_table_kr(doc,
        ['항목', '요구사항'],
        [
            ['운영체제', 'Windows 10/11 (64-bit)'],
            ['ANSYS 버전', 'ANSYS 2025 R2 (v252) 이상'],
            ['SpaceClaim', 'ANSYS SpaceClaim V252'],
            ['.NET Framework', '.NET Framework 4.7.2 이상'],
            ['설치 경로', 'C:\\ProgramData\\SpaceClaim\\AddIns\\MXDigitalTwinModeller\\V252'],
            ['디스크 공간', '약 50 MB'],
        ],
        col_widths=[4, 12]
    )

    doc.add_page_break()

    # ====================================================================
    # 3. Ribbon UI 구성
    # ====================================================================
    add_heading_kr(doc, '3. Ribbon UI 구성', level=1)
    add_paragraph_kr(doc, (
        'MX Digital Twin Modeller는 SpaceClaim 리본에 "MX Modeller" 탭을 추가합니다. '
        '이 탭은 4개의 그룹으로 구성됩니다.'
    ))

    add_table_kr(doc,
        ['그룹', '버튼 수', '기능 요약'],
        [
            ['Specimen', '4', '인장, 굽힘, 벤딩지그, 압축 시편 생성'],
            ['Advanced', '3', 'CAI, 피로, 접합부 시편 생성'],
            ['Parametric', '2', '적층 모델, Void Cut'],
            ['Mesh', '10', '재료, 메쉬, 내보내기, 단순화, 접촉, STEP, 하중, 시뮬레이션, 파이프라인, Conformal'],
        ],
        col_widths=[3, 2, 11]
    )

    add_screenshot_placeholder(doc, 'MX Modeller 리본 탭 전체 화면')

    doc.add_page_break()

    # ====================================================================
    # 4. Specimen 그룹
    # ====================================================================
    add_heading_kr(doc, '4. Specimen 그룹 (시편 생성)', level=1)
    add_paragraph_kr(doc, (
        'Specimen 그룹은 재료 시험에 필요한 시편을 파라메트릭하게 생성합니다. '
        '각 기능은 국제 규격(ASTM, ISO, DMA 등)에 맞는 치수를 자동으로 적용하며, '
        '사용자 정의도 지원합니다.'
    ))

    # ─── 4.1 인장시편 ───
    add_heading_kr(doc, '4.1 인장시편 (Tensile Specimen)', level=2)
    add_paragraph_kr(doc, (
        '인장시험 시편을 생성합니다. ASTM, ISO, DMA 등 26개 이상의 국제 규격을 지원하며, '
        '시편 형상(직사각형, 덤벨, 노치, 구멍 등)을 규격에 따라 자동으로 생성합니다.'
    ))

    add_heading_kr(doc, '지원 규격', level=3)
    add_table_kr(doc,
        ['분류', '규격', '설명'],
        [
            ['금속 인장', 'ASTM E8 Standard', '표준 금속 인장시편 (12.5mm 게이지)'],
            ['금속 인장', 'ASTM E8 SubSize', '소형 금속 인장시편 (6.4mm 게이지)'],
            ['금속 인장', 'ISO 6892-1', '국제표준 금속 인장'],
            ['플라스틱 인장', 'ASTM D638 Type I~V', '플라스틱 인장 (5가지 타입)'],
            ['플라스틱 인장', 'ISO 527-2 Type 1A/1B', '국제표준 플라스틱 인장'],
            ['노치 인장', 'ASTM E602 V/U-Notch', '노치 인장시편 (V자/U자)'],
            ['노치 인장', 'ASTM E338', '고강도 판재 노치 인장'],
            ['노치 인장', 'ASTM E292', '고온 노치 인장'],
            ['구멍 시편', 'ASTM D5766 (OHT)', '오픈홀 인장'],
            ['구멍 시편', 'ASTM D6484 (OHC)', '오픈홀 압축'],
            ['구멍 시편', 'ASTM D6742 (FHT)', '필드홀 인장'],
            ['구멍 시편', 'ASTM D5961 (Bearing)', '볼트 지지력 시험'],
            ['복합재 인장', 'ASTM D3039', '복합재 인장'],
            ['PCB', 'IPC-TM-650 Tensile', 'PCB 인장'],
            ['PCB', 'IPC-TM-650 PTH Pull', 'PCB 관통홀 인발'],
            ['DMA 인장', 'Rectangle / DogBone', 'DMA 인장시편'],
            ['전단', 'ASTM D5379 (Iosipescu)', 'V노치 전단'],
            ['전단', 'ASTM D7078 (V-Notch Rail)', 'V노치 레일전단'],
            ['사용자 정의', 'Custom', '사용자 정의 치수'],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('기본 치수: ', 'Gauge Length, Gauge Width, Thickness, Total Length'),
        ('그립 영역: ', 'Grip Width, Grip Length, Fillet Radius'),
        ('노치 옵션: ', 'Notch Depth, Angle (V-notch), Radius (U-notch), Double Notch'),
        ('구멍 옵션: ', 'Hole Diameter, Elliptical Hole (Major/Minor Axis)'),
        ('탭 옵션: ', 'Tab Length, Tab Thickness, Rectangular Mode'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        'Modeless 대화창: SpaceClaim 3D 뷰어와 동시 조작 가능',
        '실시간 미리보기: Preview 버튼으로 생성 전 형상 확인',
        '규격별 자동 기본값: 규격 선택 시 치수 자동 적용',
        '동적 UI: 선택한 시편 유형에 따라 관련 옵션만 표시',
        '그립 지그 자동 생성: 시편과 함께 그립 부자재 자동 생성',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, '인장시편 대화창 및 생성 결과')

    doc.add_page_break()

    # ─── 4.2 굽힘시편 ───
    add_heading_kr(doc, '4.2 굽힘시편 (Bending Specimen)', level=2)
    add_paragraph_kr(doc, (
        '3점 굽힘(3-Point Bending) 및 4점 굽힘(4-Point Bending) 시편을 생성합니다. '
        '지지대(Support)와 가압봉(Loading Nose)을 포함한 완전한 시험 셋업을 자동으로 구성합니다.'
    ))

    add_heading_kr(doc, '지원 규격', level=3)
    add_table_kr(doc,
        ['시험 유형', '규격', '설명'],
        [
            ['3점 굽힘', 'ASTM D790 (16:1)', '플라스틱 굽힘 (Span/Thickness=16)'],
            ['3점 굽힘', 'ASTM D790 (32:1)', '플라스틱 굽힘 (Span/Thickness=32)'],
            ['3점 굽힘', 'ISO 178', '국제표준 플라스틱 굽힘'],
            ['3점 굽힘', 'ASTM D7264 (32:1)', '복합재 굽힘'],
            ['3점 굽힘', 'ASTM C1161 (Config A/B)', '세라믹 굽힘'],
            ['3점 굽힘', 'DMA Standard', 'DMA 3점 굽힘'],
            ['4점 굽힘', 'ASTM D6272 (16:1)', '플라스틱 4점 굽힘'],
            ['4점 굽힘', 'ASTM D6272 (32:1)', '플라스틱 4점 굽힘'],
            ['4점 굽힘', 'ASTM C1161 (Config B/C)', '세라믹 4점 굽힘'],
            ['4점 굽힘', 'DMA Standard', 'DMA 4점 굽힘'],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('시편 치수: ', 'Length, Width, Thickness'),
        ('3점 굽힘: ', 'Span (Span/Thickness 비율 자동 표시)'),
        ('4점 굽힘: ', 'Outer Span, Inner Span (비율 자동 표시)'),
        ('지그 치수: ', 'Support Diameter/Height, Loading Nose Diameter/Height'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        '3점/4점 통합 대화창: 단일 인터페이스에서 모드 전환',
        '규격 준수 실시간 경고: 비율/치수 범위 초과 시 경고 표시',
        '지지대 + 가압봉 자동 생성: 시편과 함께 시험 지그 생성',
        '실시간 미리보기: Preview 기능으로 생성 전 확인',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, '굽힘시편 대화창 및 3점/4점 굽힘 결과')

    doc.add_page_break()

    # ─── 4.3 벤딩 지그 ───
    add_heading_kr(doc, '4.3 벤딩 지그 (Bending Fixture)', level=2)
    add_paragraph_kr(doc, (
        '기존에 생성된 바디에 3점 벤딩 지지구조를 적용합니다. '
        '바디의 바운딩박스를 분석하여 자동으로 방향을 감지하고, '
        '지지대와 가압봉을 정확한 위치에 생성합니다.'
    ))

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('바디 선택: ', '드롭다운에서 기존 바디 선택 또는 사전 선택(Pre-selection) 지원'),
        ('방향 할당: ', 'X/Y/Z축을 Span/Width/Loading 방향에 할당'),
        ('스팬 모드: ', 'Ratio (바디 길이의 %) 또는 Absolute (mm 직접 입력)'),
        ('지그 치수: ', 'Support/Loading Nose 직경 및 높이'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        'Auto-Detect: 바디 형상으로부터 Span/Width/Loading 방향 자동 감지',
        '실시간 치수 표시: 바디 크기(mm) 표시',
        '충돌 방지: 동일 축 중복 할당 시 자동 스왑',
        '3개 지그 바디 생성: 지지대 2개 + 가압봉 1개',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, '벤딩 지그 적용 대화창 및 결과')

    doc.add_page_break()

    # ─── 4.4 압축시편 ───
    add_heading_kr(doc, '4.4 압축시편 (Compression Specimen)', level=2)
    add_paragraph_kr(doc, (
        '압축 시험 시편을 생성합니다. 각형(Prism)과 원통형(Cylinder) 형상을 지원하며, '
        '선택적으로 압축 플래튼을 함께 생성할 수 있습니다.'
    ))

    add_heading_kr(doc, '지원 규격', level=3)
    add_table_kr(doc,
        ['규격', '형상', '기본 치수'],
        [
            ['ASTM D695 - Prism', '각형', '12.7 x 12.7 x 25.4 mm'],
            ['ASTM D695 - Cylinder', '원통형', 'Ø12.7 x 25.4 mm'],
            ['ISO 604 - Modulus', '각형', '50 x 10 x 4 mm'],
            ['ISO 604 - Strength', '각형', '10 x 10 x 4 mm'],
            ['ASTM E9 - Short (L/D=2)', '원통형', 'Ø12.7 x 25.4 mm'],
            ['ASTM E9 - Medium (L/D=3)', '원통형', 'Ø12.7 x 38.1 mm'],
            ['Custom', '사용자 선택', '사용자 정의'],
        ],
        col_widths=[5, 3, 8]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('형상 선택: ', 'Prism (Width, Depth, Height) / Cylinder (Diameter, Height)'),
        ('플래튼 옵션: ', '생성 여부, 플래튼 직경/높이'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_screenshot_placeholder(doc, '압축시편 대화창 및 각형/원통형 결과')

    doc.add_page_break()

    # ====================================================================
    # 5. Advanced 그룹
    # ====================================================================
    add_heading_kr(doc, '5. Advanced 그룹 (고급 시험)', level=1)
    add_paragraph_kr(doc, (
        'Advanced 그룹은 특수 시험 규격에 맞는 고급 시편을 생성합니다. '
        'CAI(충격 후 압축), 피로, 접합부 시편을 포함합니다.'
    ))

    # ─── 5.1 CAI ───
    add_heading_kr(doc, '5.1 CAI시편 (Compression After Impact)', level=2)
    add_paragraph_kr(doc, (
        'CAI (Compression After Impact) 시험용 시편을 생성합니다. '
        '충격 시험 후 잔류 압축 강도를 측정하기 위한 패널형 시편으로, '
        '선택적으로 지지 지그와 손상 영역을 함께 모델링합니다.'
    ))

    add_heading_kr(doc, '지원 규격', level=3)
    add_table_kr(doc,
        ['규격', '설명', '기본 치수'],
        [
            ['ASTM D7137/D7136', 'CAI 표준', '150 x 100 mm, 4 mm 두께'],
            ['ASTM D6264', '낙추충격', '150 x 100 mm, 5 mm 두께'],
            ['Boeing BSS 7260', '항공용 CAI', '152.4 x 101.6 mm (6x4 inch)'],
            ['Custom', '사용자 정의', '사용자 정의 치수'],
        ],
        col_widths=[5, 4, 7]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('패널 치수: ', 'Panel Length, Panel Width, Thickness'),
        ('지그 옵션: ', '지그 생성 여부, 지그 두께, 창(Window) 크기, Clearance'),
        ('손상 영역: ', '생성 여부, 원형/타원형, 직경(또는 장축/단축), 깊이(%)'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        '지지 지그 자동 생성: 시편 + 지그 + 손상 영역 일괄 생성',
        '손상 형상 선택: 원형(Circular) 또는 타원형(Elliptical)',
        '깊이 비율 제어: 손상 깊이를 두께 백분율로 지정',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'CAI시편 대화창 및 생성 결과 (시편+지그+손상영역)')

    doc.add_page_break()

    # ─── 5.2 피로시편 ───
    add_heading_kr(doc, '5.2 피로시편 (Fatigue Specimen)', level=2)
    add_paragraph_kr(doc, (
        '피로 시험 시편을 생성합니다. 하중제어(HCF), 변형률제어(LCF), '
        '피로균열성장(FCG), 다축 피로 등 다양한 피로 시험 규격을 지원합니다.'
    ))

    add_heading_kr(doc, '지원 규격', level=3)
    add_table_kr(doc,
        ['규격', '시험 유형', '설명'],
        [
            ['ASTM E466 (Uniform)', '하중제어 HCF', 'Dog-bone 형태 평판 시편, 큰 필렛 반경'],
            ['ASTM E466 (Hourglass)', '하중제어 HCF', '모래시계형, 연속 곡률로 최소 단면 중앙'],
            ['ASTM E606', '변형률제어 LCF', '원형 단면, 짧은 게이지, 저주기 피로'],
            ['ASTM E647 CT', '피로균열성장', 'Compact Tension, 핀홀+노치'],
            ['ASTM E647 M(T)', '피로균열성장', 'Middle Tension, 중앙 양방향 슬롯'],
            ['ASTM E2207', '다축 피로', '박벽 원통형, 축방향+비틀림'],
            ['Custom', '사용자 정의', '사용자 정의 치수'],
        ],
        col_widths=[4, 3.5, 8.5]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('기본 치수 (E466/E606): ', 'Gauge Length, Gauge Width, Thickness, Gauge Diameter, Grip Width/Length, Total Length, Fillet Radius, Hourglass Radius'),
        ('CT 파라미터 (E647): ', 'CT Width, CT Thickness, Initial Crack Length, Pin Hole Diameter, Notch Width'),
        ('M(T) 파라미터 (E647): ', 'MT Width, MT Length, MT Thickness, Slot Half Length, Slot Width'),
        ('튜브 파라미터 (E2207): ', 'Outer/Inner Diameter, Gauge Length, Total Length, Grip OD'),
        ('옵션: ', 'Create Grips (그립 부자재 생성 여부)'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        '동적 UI: 규격에 따라 관련 파라미터 패널만 표시',
        '6가지 시편 유형: 평판 → 원형 → CT → M(T) → 원통형',
        '그립 부자재 선택적 생성',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, '피로시편 대화창 및 각 유형별 생성 결과')

    doc.add_page_break()

    # ─── 5.3 접합부시편 ───
    add_heading_kr(doc, '5.3 접합부시편 (Joint Specimen)', level=2)
    add_paragraph_kr(doc, (
        '접합부 시험 시편을 생성합니다. 접착 접합의 강도 평가를 위한 '
        '다양한 접합 형태(겹치기, 스카프, 맞대기, T형)를 지원합니다.'
    ))

    add_heading_kr(doc, '지원 규격', level=3)
    add_table_kr(doc,
        ['규격/유형', '설명'],
        [
            ['ASTM D1002 - Single Lap Shear', '단일겹치기 전단접합 (25.4mm 오버랩)'],
            ['ASTM D3528 - Double Lap Shear', '이중겹치기 전단접합 (양면 접착, 편심 감소)'],
            ['Scarf Joint', '스카프 접합 (경사면 접착, 균일 응력 분포)'],
            ['Butt Joint', '맞대기 접합 (끝단 맞대기, 인장 강도 측정)'],
            ['T-Joint Pull-off', 'T형 접합 (플랜지+웹 구조, 박리 시험)'],
            ['Custom', '사용자 정의 치수'],
        ],
        col_widths=[6, 10]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('피착재 치수: ', 'Adherend Width, Length, Thickness'),
        ('접합 파라미터: ', 'Overlap Length, Adhesive Thickness, Scarf Angle'),
        ('T-Joint: ', 'Flange Length, Web Height, Web Thickness, Fillet Bond Size'),
        ('옵션: ', 'Create Adhesive Body (접착층 별도 생성), Create Grips (그립 생성)'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        '동적 UI: 접합 유형에 따라 관련 파라미터만 표시',
        '접착층 별도 바디 생성 옵션 (접촉 해석용)',
        'T-Joint 전용 패널: 플랜지/웹/필렛 치수 제어',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, '접합부시편 대화창 및 각 유형별 생성 결과')

    doc.add_page_break()

    # ====================================================================
    # 6. Parametric 그룹
    # ====================================================================
    add_heading_kr(doc, '6. Parametric 그룹 (파라메트릭 모델러)', level=1)
    add_paragraph_kr(doc, (
        'Parametric 그룹은 적층 복합재 모델링과 Boolean 연산 기반 Void Cut 기능을 제공합니다.'
    ))

    # ─── 6.1 적층 모델 ───
    add_heading_kr(doc, '6.1 적층 모델 (Laminate)', level=2)
    add_paragraph_kr(doc, (
        '적층 복합재 모델을 생성합니다. 3가지 생성 모드를 지원하며, '
        '레이어별 두께와 이름을 개별 제어할 수 있습니다.'
    ))

    add_heading_kr(doc, '생성 모드', level=3)
    add_table_kr(doc,
        ['모드', '입력', '설명'],
        [
            ['Rectangular', 'Width, Length', '직사각형 적층판 생성, 적층 방향(Z+/Z-/Y+/Y-) 선택'],
            ['Surface', 'Face 선택', '3D 뷰에서 면 선택 후 Offset 방향으로 레이어 생성'],
            ['Solid', 'Body 선택', '기존 솔리드의 두께/방향 자동 감지, 레이어 분할'],
        ],
        col_widths=[3, 4, 9]
    )

    add_heading_kr(doc, '레이어 관리', level=3)
    params = [
        ('DataGridView: ', '레이어 목록 (이름, 두께 mm)'),
        ('추가/삭제: ', 'Add/Remove 버튼으로 레이어 추가/삭제'),
        ('순서 변경: ', 'Move Up/Move Down 버튼으로 적층 순서 조정'),
        ('총 두께: ', '전체 두께 및 레이어 수 실시간 표시'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '옵션', level=3)
    opts = [
        'Share Topology: 레이어 간 공유 위상 생성 (Rectangular 모드)',
        'Interface NS: 레이어 계면 Named Selection 자동 생성',
        'Match Thickness: 솔리드 두께에 맞게 레이어 두께 자동 분배 (Solid 모드)',
        'Delete Original: 원본 바디 삭제 (Solid 모드)',
    ]
    for o in opts:
        add_bullet_kr(doc, o)

    add_screenshot_placeholder(doc, '적층 모델 대화창 (Rectangular/Surface/Solid 모드)')

    doc.add_page_break()

    # ─── 6.2 Void Cut ───
    add_heading_kr(doc, '6.2 Void Cut (보이드 컷)', level=2)
    add_paragraph_kr(doc, (
        '솔리드 모델에서 빈 공간(Void)을 모델링합니다. '
        '커터 볼륨을 정의하여 교차하는 모든 바디에 Boolean 연산을 수행합니다.'
    ))

    add_heading_kr(doc, '커터 정의 모드', level=3)
    add_table_kr(doc,
        ['모드', '설명'],
        [
            ['기존 바디 선택', 'CheckedListBox에서 커터 바디 및 대상 바디를 각각 선택'],
            ['형상 직접 정의', 'Cuboid (W/H/D), Cylinder (R/H), Sphere (R) 중 선택하여 치수 입력'],
        ],
        col_widths=[4, 12]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('형상: ', 'Cuboid (Width/Height/Depth), Cylinder (Radius/Height), Sphere (Radius)'),
        ('위치: ', 'X, Y, Z (mm, 중심 기준)'),
        ('Boolean 연산: ', 'Subtract, Unite, Intersect'),
        ('Offset: ', '커터 볼륨 확장 (mm)'),
        ('커터 삭제: ', '연산 후 커터 바디 삭제 여부'),
        ('NS 생성: ', '절단 영역에 Named Selection 자동 생성'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '고급 기능', level=3)
    adv = [
        '미리보기 모드: 교차 후보 바디 사전 확인',
        '배치 모드: CSV 파일에서 다중 Void 정의 로드',
        '결과 로그: 성공/스킵/실패 개수 표시',
        '볼륨 정보: 각 바디의 볼륨 및 바운딩박스 표시',
    ]
    for a in adv:
        add_bullet_kr(doc, a)

    add_screenshot_placeholder(doc, 'Void Cut 대화창 및 실행 결과')

    doc.add_page_break()

    # ====================================================================
    # 7. Mesh 그룹
    # ====================================================================
    add_heading_kr(doc, '7. Mesh 그룹 (메쉬 및 해석 설정)', level=1)
    add_paragraph_kr(doc, (
        'Mesh 그룹은 재료 물성 적용, 메쉬 생성, 접촉면 감지, 내보내기, '
        '시뮬레이션 설정 등 해석 준비의 전체 워크플로를 지원합니다.'
    ))

    # ─── 7.1 Material ───
    add_heading_kr(doc, '7.1 Material (재료 물성)', level=2)
    add_paragraph_kr(doc, (
        '바디에 재료 물성을 적용합니다. LS-DYNA 호환 단위계(mm-tonne-s)를 사용하며, '
        '프리셋(Steel, Aluminum, CFRP) 또는 사용자 정의 물성을 입력할 수 있습니다.'
    ))

    add_heading_kr(doc, '물성 항목', level=3)
    add_table_kr(doc,
        ['기호', '물성', '단위'],
        [
            ['ρ', 'Density (밀도)', 'tonne/mm³'],
            ['E', "Young's Modulus (영률)", 'MPa'],
            ['ν', "Poisson's Ratio (포아송비)", '-'],
            ['G', 'Shear Modulus (전단 탄성률)', 'MPa'],
            ['σ_u', 'Ultimate Stress (극한 응력)', 'MPa'],
            ['α', 'CTE (열팽창계수)', '1/°C'],
            ['k', 'Thermal Conductivity (열전도율)', 'W/(mm·°C)'],
            ['c_p', 'Specific Heat (비열)', 'J/(kg·°C)'],
        ],
        col_widths=[2, 6, 4]
    )

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        '좌측: 전체 바디 물성 요약 테이블',
        '우측: 물성 입력 컨트롤',
        '키워드 필터: 바디 이름으로 일괄 적용',
        'SpaceClaim 바디 속성(Attributes)에 물성 저장',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Material 대화창')

    # ─── 7.2 Mesh 설정 ───
    add_heading_kr(doc, '7.2 Mesh 설정', level=2)
    add_paragraph_kr(doc, (
        '바디별 메쉬 크기를 설정하고 격자를 생성합니다. '
        '방향별(X/Y/Z) 크기 제어, 요소 형상 선택, 일괄 업데이트를 지원합니다.'
    ))

    add_heading_kr(doc, '바디별 설정 항목', level=3)
    add_table_kr(doc,
        ['파라미터', '설명', '옵션'],
        [
            ['Directional Sizing', 'X, Y, Z 방향 메쉬 크기', 'mm 단위'],
            ['Element Shape', '요소 형상', 'Tet, Hex, Quad, Tri'],
            ['Midside Nodes', '중간절점', 'Dropped, Kept, Auto'],
            ['Growth Rate', '성장률', '1.0 ~ 5.0'],
            ['Size Function', '크기 함수', 'Curv+Prox, Curv, Prox, Fixed'],
        ],
        col_widths=[4, 5, 7]
    )

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        'DataGridView: 모든 바디 목록 및 개별 설정 편집',
        '키워드 필터: 바디 이름으로 일괄 선택 및 설정',
        '변경 감지: 수정된 바디만 재메쉬',
        '동일 설정 그룹화: 같은 설정의 바디를 배치 메쉬',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Mesh 설정 대화창')

    # ─── 7.3 Mesh Export ───
    add_heading_kr(doc, '7.3 Mesh Export (메쉬 내보내기)', level=2)
    add_paragraph_kr(doc, (
        '생성된 메쉬를 다양한 해석 소프트웨어 형식으로 내보냅니다.'
    ))

    add_heading_kr(doc, '지원 형식', level=3)
    add_table_kr(doc,
        ['형식', '확장자', '비고'],
        [
            ['LS-DYNA', '.k', '재료 물성 패칭, 제어 카드 자동 삽입'],
            ['ANSYS Mechanical', '.cdb', '-'],
            ['Abaqus', '.inp', '-'],
            ['Fluent Mesh', '.msh', '-'],
            ['CGNS', '.cgns', '-'],
        ],
        col_widths=[4, 3, 9]
    )

    add_heading_kr(doc, 'LS-DYNA 후처리 (자동)', level=3)
    feat_list = [
        '재료 물성 패칭: 바디 속성에서 읽어 자동 대체',
        '제어 카드 삽입: *CONTROL_* 카드 자동 생성',
        '하중 곡선 삽입: Named Selection 기반 하중 곡선',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Mesh Export 대화창')

    doc.add_page_break()

    # ─── 7.4 Simplify ───
    add_heading_kr(doc, '7.4 Simplify (바디 단순화)', level=2)
    add_paragraph_kr(doc, (
        '키워드 매칭 바디를 BoundingBox 또는 Shell로 단순화합니다. '
        '해석에 중요하지 않은 부품(체결재, 지그 등)의 메쉬 수를 줄이기 위해 사용합니다.'
    ))

    add_heading_kr(doc, '단순화 모드', level=3)
    add_table_kr(doc,
        ['모드', '설명'],
        [
            ['BoundingBox', '바디를 최소 바운딩 박스로 대체'],
            ['SolidToShell', '솔리드를 셸(표면)로 변환'],
        ],
        col_widths=[4, 12]
    )

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        'DataGridView로 규칙 정의 (키워드 + 모드)',
        '규칙 추가/삭제 버튼',
        'CSV 가져오기/내보내기',
        '일괄 실행 및 결과 로그',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Simplify 대화창')

    # ─── 7.5 Contact ───
    add_heading_kr(doc, '7.5 Contact (접촉면 감지)', level=2)
    add_paragraph_kr(doc, (
        '바디 간 접촉면을 자동으로 감지하여 Named Selection을 생성합니다. '
        '키워드 기반 필터링과 수동 추가를 모두 지원합니다.'
    ))

    add_heading_kr(doc, '감지 모드', level=3)
    add_table_kr(doc,
        ['모드', '조건', '설명'],
        [
            ['전체 모드', '키워드 비어있음', '모든 바디 간 접촉 감지'],
            ['단일 모드', 'Keyword A만 입력', 'Keyword A 바디 ↔ 나머지 바디'],
            ['쌍 모드', 'Keyword A + B', 'Keyword A 바디 ↔ Keyword B 바디'],
        ],
        col_widths=[3, 4, 9]
    )

    add_heading_kr(doc, '접촉 유형', level=3)
    add_table_kr(doc,
        ['유형', '설명'],
        [
            ['면 (Face)', '면-면 접촉'],
            ['에지 (Edge)', '에지 접촉'],
            ['원통 (Cylinder)', '원통-원통 접촉'],
            ['평-원 (Plane-Cylinder)', '평면-원통 접촉'],
        ],
        col_widths=[4, 12]
    )

    add_heading_kr(doc, '고급 기능', level=3)
    feat_list = [
        '결과 테이블: 유형, Body A/B, 면적, Prefix, 상태 표시',
        '3D 동기화: 행 선택 시 3D 뷰에서 해당 면 하이라이트',
        '수동 추가: 3D 뷰에서 2개 면 선택하여 접촉 쌍 추가',
        'Prefix 관리: 키워드 기반 자동 할당 또는 일괄 적용',
        'CSV 가져오기/내보내기: 접촉 설정 저장/복원',
        'NS 생성/삭제: 선택한 접촉 쌍에 대한 Named Selection 생성/삭제',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Contact 대화창 및 접촉면 감지 결과')

    doc.add_page_break()

    # ─── 7.6 STEP Export ───
    add_heading_kr(doc, '7.6 STEP Export', level=2)
    add_paragraph_kr(doc, (
        '현재 모델을 STEP (.stp) 파일로 내보냅니다. '
        '다른 CAD/CAE 소프트웨어와의 데이터 교환에 사용됩니다.'
    ))
    add_screenshot_placeholder(doc, 'STEP Export 실행')

    # ─── 7.7 Load ───
    add_heading_kr(doc, '7.7 Load (하중 정의)', level=2)
    add_paragraph_kr(doc, (
        'Named Selection에 시간이력 하중을 정의합니다. '
        '3가지 입력 모드를 지원하며, 시간 이력 및 FFT 그래프를 실시간으로 표시합니다.'
    ))

    add_heading_kr(doc, '입력 모드', level=3)
    add_table_kr(doc,
        ['모드', '설명'],
        [
            ['Expression', '수학식으로 하중 함수 정의 (sin, cos, exp 등)'],
            ['Tabular', '시간-하중 데이터를 직접 입력 또는 붙여넣기'],
            ['PWM', '펄스폭변조 (Carrier Freq, Output Amp, Harmonics)'],
        ],
        col_widths=[3, 13]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('공통: ', 'Named Selection 선택, Load Name, End Time, dt'),
        ('Expression: ', '수식 입력 (수식 도움말 표시)'),
        ('PWM: ', 'Carrier Frequency, Output Amplitude, Bipolar, Target Frequency'),
        ('PWM 고조파: ', 'Harmonics 테이블 (추가/삭제), Optimize 버튼'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '그래프 기능', level=3)
    feat_list = [
        '시간 이력 그래프: 생성된 하중 파형 실시간 표시',
        'FFT 그래프: 주파수 영역 분석 결과 표시',
        'X축 줌: 드래그로 특정 시간/주파수 구간 확대',
        '하중 목록: 정의된 하중 추가/수정/삭제 관리',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Load 대화창 (Expression/PWM 모드)')

    # ─── 7.8 Simulation Setup ───
    add_heading_kr(doc, '7.8 Simulation Setup (시뮬레이션 설정)', level=2)
    add_paragraph_kr(doc, (
        'LS-DYNA 시뮬레이션 조건을 설정합니다. '
        '모드 해석(Modal Analysis) 설정을 지원하며, 제어 카드를 미리보기하고 내보낼 수 있습니다.'
    ))

    add_heading_kr(doc, '설정 항목', level=3)
    add_table_kr(doc,
        ['카테고리', '파라미터'],
        [
            ['일반', 'Analysis Type (Modal), Title'],
            ['고유값', 'Number of Modes, Min/Max Frequency, Eigen Method'],
            ['솔버', 'Solver, Auto SPC, Negative Eigenvalue, Geometric Stiffness, IM Form'],
            ['출력', 'EIGOUT, D3PLOT, NODOUT, ELOUT'],
            ['추가', 'Energy Balance, Hourglass Control, Accuracy'],
        ],
        col_widths=[3, 13]
    )

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        '제어 카드 미리보기: 생성될 LS-DYNA 키워드 실시간 표시',
        '클립보드 복사: 제어 카드를 클립보드에 복사',
        '파일 내보내기: 제어 카드를 별도 파일로 저장',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Simulation Setup 대화창')

    doc.add_page_break()

    # ─── 7.9 Pipeline ───
    add_heading_kr(doc, '7.9 Pipeline (일괄 실행)', level=2)
    add_paragraph_kr(doc, (
        '단순화 → 재료 적용 → 접촉 감지 → 메쉬 생성 → 내보내기의 '
        '전체 워크플로를 순차적으로 일괄 실행합니다.'
    ))

    add_heading_kr(doc, '파이프라인 단계', level=3)
    add_table_kr(doc,
        ['단계', '기능', '주요 설정'],
        [
            ['1', '단순화 (Simplify)', 'Keyword, Mode (BoundingBox/SolidToShell)'],
            ['2', '재료 적용 (Material)', 'Preset (Steel/Al/CFRP), Keyword 필터'],
            ['3', '접촉 감지 (Contact)', 'Keyword A/B, Tolerance (mm)'],
            ['4', '메쉬 생성 (Mesh)', 'Size, Shape, Midside, Growth Rate, Size Function'],
            ['5', '내보내기 (Export)', 'Format (LS-DYNA/ANSYS/Abaqus/Fluent/CGNS), 경로'],
        ],
        col_widths=[1, 4, 11]
    )

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        '각 단계 선택적 실행: 필요한 단계만 체크하여 실행',
        'Continue on Error: 한 단계 실패 시에도 다음 단계 계속 실행 옵션',
        '진행 표시: 단계별 진행률 바 및 타임스탬프 로그',
        '설정 저장/불러오기: .cfg 파일로 파이프라인 설정 저장/복원',
        '최종 요약: 성공/실패 단계 수 표시',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Pipeline 대화창')

    # ─── 7.10 Conformal Mesh ───
    add_heading_kr(doc, '7.10 Conformal Mesh (적합 메쉬)', level=2)
    add_paragraph_kr(doc, (
        'STEP 어셈블리로부터 공유 위상(Share Topology) 기반의 적합 메쉬를 생성합니다. '
        '파트 간 계면에서 절점이 일치하는 메쉬를 자동으로 생성합니다.'
    ))

    add_heading_kr(doc, '워크플로', level=3)
    steps = [
        'STEP 파일 로드 (현재 파트 또는 파일 열기)',
        '계면 검출 (Tolerance 설정 후 Detect 실행)',
        '계면 목록 확인 (DataGridView에서 검출된 계면 확인)',
        'Share Topology 적용',
        '메쉬 생성 (Element Size, Growth Rate, Strategy 설정)',
        '메쉬 내보내기',
    ]
    for i, s in enumerate(steps, 1):
        add_bullet_kr(doc, f'{s}', bold_prefix=f'{i}. ')

    add_heading_kr(doc, '메쉬 설정', level=3)
    params = [
        ('Element Size: ', '전체 요소 크기 (mm)'),
        ('Growth Rate: ', '성장률'),
        ('Strategy: ', '메쉬 전략 선택'),
        ('Curvature & Proximity: ', '곡률 및 근접성 기반 크기 조정'),
        ('Cylinder Split: ', '원통면 분할 및 분할 수 설정'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_screenshot_placeholder(doc, 'Conformal Mesh 대화창')

    doc.add_page_break()

    # ====================================================================
    # 기능 요약 테이블
    # ====================================================================
    add_heading_kr(doc, '기능 요약', level=1)
    add_table_kr(doc,
        ['No.', '그룹', '기능명', '설명'],
        [
            ['1', 'Specimen', '인장시편', '26+ 규격 인장시편 생성'],
            ['2', 'Specimen', '굽힘시편', '3점/4점 굽힘시편 + 지그 생성'],
            ['3', 'Specimen', '벤딩 지그', '기존 바디에 벤딩 지지구조 적용'],
            ['4', 'Specimen', '압축시편', '각형/원통형 압축시편 생성'],
            ['5', 'Advanced', 'CAI시편', 'CAI 패널 + 지그 + 손상영역'],
            ['6', 'Advanced', '피로시편', '6종 피로시험 규격 지원'],
            ['7', 'Advanced', '접합부시편', '5종 접합 유형 지원'],
            ['8', 'Parametric', '적층 모델', '3모드 적층 복합재 모델링'],
            ['9', 'Parametric', 'Void Cut', 'Boolean 연산 기반 보이드 생성'],
            ['10', 'Mesh', 'Material', '재료 물성 적용 (8개 항목)'],
            ['11', 'Mesh', 'Mesh 설정', '바디별 방향성 메쉬 크기 제어'],
            ['12', 'Mesh', 'Mesh Export', '5가지 형식 메쉬 내보내기'],
            ['13', 'Mesh', 'Simplify', '바디 단순화 (BBox/Shell)'],
            ['14', 'Mesh', 'Contact', '접촉면 자동 감지 + NS 생성'],
            ['15', 'Mesh', 'STEP Export', 'STEP 파일 내보내기'],
            ['16', 'Mesh', 'Load', '시간이력 하중 정의 (3모드)'],
            ['17', 'Mesh', 'Simulation', 'LS-DYNA 모드해석 설정'],
            ['18', 'Mesh', 'Pipeline', '전체 워크플로 일괄 실행'],
            ['19', 'Mesh', 'Conformal', 'STEP 기반 적합 메쉬'],
        ],
        col_widths=[1, 2.5, 3.5, 9]
    )

    # ====================================================================
    # 저장
    # ====================================================================
    output_path = os.path.join(os.path.dirname(__file__), 'MXDigitalTwinModeller_SpaceClaim_Report.docx')
    doc.save(output_path)
    print(f'Report saved: {output_path}')
    return output_path


if __name__ == '__main__':
    create_report()
