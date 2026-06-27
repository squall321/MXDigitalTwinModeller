from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# -- style --
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def add_heading_kr(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Malgun Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return h

def add_table_with_style(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Malgun Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2F5496', qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Malgun Gothic'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return table

# ============================
# Title
# ============================
add_heading_kr('LS-DYNA 물성 분석 보고서', level=0)
doc.add_paragraph('모바일 기기 낙하(Drop) 시뮬레이션 모델')
doc.add_paragraph('')

# -- 1. Overview --
add_heading_kr('1. 모델 개요', level=1)
p = doc.add_paragraph()
p.add_run('LS-PrePost 2025 R1').bold = True
p.add_run('으로 생성된 LS-DYNA 키워드 파일 분석.')
doc.add_paragraph('  - 종료 시간: 0.001 sec (1 ms) - 낙하 충격 해석')
doc.add_paragraph('  - 질량 스케일링: dt2ms = -2.6E-9 sec')
doc.add_paragraph('  - 접촉: AUTOMATIC_GENERAL + TIED_SURFACE_TO_SURFACE_OFFSET')
doc.add_paragraph('  - 총 물성 수: 22개')
doc.add_paragraph('  - 참고: 에폭시 계열에 고무 성분이 혼합된 복합 재질 포함')
doc.add_paragraph('')

# -- 2. Material List --
add_heading_kr('2. 전체 물성 목록 (22개)', level=1)

add_heading_kr('2.1 기본 탄성 물성 - *MAT_ELASTIC (21개)', level=2)
doc.add_paragraph('밀도(ro), 탄성계수(E), 푸아송비(PR) 3개 파라미터만 정의됨.')

elastic_rows = [
    ['1',  'BGA40G',        '4.00E-9', '40,000',  '0.30', '솔더볼 (BGA)'],
    ['2',  'PC-GF30%',      '1.41E-9', '8,150',   '0.44', 'PC+유리섬유 30%'],
    ['3',  'Mobile.STS304', '7.90E-9', '190,000', '0.30', 'STS304 스테인리스강'],
    ['4',  'Ground',        '7.83E-6', '207,000', '0.288','바닥면 (강체)'],
    ['6',  'SOLDER_50G',    '7.40E-9', '50,000',  '0.44', '솔더'],
    ['7',  'Epoxy_1000M',   '1.30E-9', '1,000',   '0.37', '에폭시 (고무 혼합 가능)'],
    ['8',  'PCB_30G',       '4.34E-9', '30,000',  '0.275','PCB (FR-4)'],
    ['9',  'Panel',         '2.23E-9', '7,500',   '0.30', '디스플레이 패널'],
    ['10', 'AL7003H',       '2.70E-9', '63,100',  '0.33', '알루미늄 7003-H'],
    ['11', 'Epoxy_100M',    '1.30E-9', '100',     '0.37', '에폭시 (저강성, 고무 혼합)'],
    ['12', 'TAPE_1620MPa',  '1.10E-9', '1,620',   '0.37', '접착 테이프'],
    ['13', 'PC_QF1035',     '1.18E-9', '2,344',   '0.41', '폴리카보네이트'],
    ['14', 'GG8(79G)',      '2.41E-9', '79,000',  '0.21', '고릴라 글래스 8'],
    ['15', 'Mobile.Epoxy',  '1.30E-9', '3,000',   '0.37', '에폭시 (고무 혼합 가능)'],
    ['16', 'PC-GF50%',      '1.60E-9', '12,000',  '0.43', 'PC+유리섬유 50%'],
    ['17', 'PC-GF10%',      '1.29E-9', '2,600',   '0.44', 'PC+유리섬유 10%'],
    ['18', 'PC_BATT',       '2.36E-9', '2,500',   '0.30', '배터리 케이스 (PC)'],
    ['19', 'PCGF30_LS3302', '1.40E-9', '4,880',   '0.37', 'PC+GF30% (LS3302)'],
    ['20', 'Cu',            '8.92E-9', '114,000', '0.31', '구리'],
    ['21', 'SUZY_1520M',    '1.30E-9', '1,520',   '0.40', '실리콘/접착제'],
    ['24', 'Panel1187',     '2.23E-9', '1,187',   '0.30', '디스플레이 패널 (얇은)'],
]
add_table_with_style(
    ['MID', '재료명', 'ro (tonne/mm3)', 'E (MPa)', 'PR', '추정 재질'],
    elastic_rows
)

doc.add_paragraph('')
add_heading_kr('2.2 소성 물성 - *MAT_PIECEWISE_LINEAR_PLASTICITY (1개)', level=2)

plastic_rows = [
    ['23', 'C7521_1_2H', '8.70E-9', '124,000', '0.28', '310-440', '인청동 합금'],
]
add_table_with_style(
    ['MID', '재료명', 'ro', 'E (MPa)', 'PR', 'Sy-Su (MPa)', '추정 재질'],
    plastic_rows
)
doc.add_paragraph('  - eps=0.00 -> 310 MPa / eps=0.23 -> 440 MPa / eps=1.00 -> 440 MPa (포화)')

# -- 3. Advanced Analysis --
doc.add_paragraph('')
add_heading_kr('3. 물성 고도화 분석', level=1)

add_heading_kr('3.1 점탄성/초탄성 모델 필요 - 고분자/솔더/접착제 (7개)', level=2)
doc.add_paragraph('에폭시 계열은 고무 성분이 혼합되어 있어 점탄성 거동이 더욱 뚜렷함.')

ve_rows = [
    ['1',  'BGA40G',       '솔더볼',           'MAT_024/MAT_106', '크리프+변형률속도 의존'],
    ['6',  'SOLDER_50G',   '솔더',             'MAT_024/MAT_106', '크리프+변형률속도 의존'],
    ['7',  'Epoxy_1000M',  '에폭시+고무 혼합',  'MAT_076',         '점탄성 (Prony series)'],
    ['11', 'Epoxy_100M',   '에폭시+고무 (저E)', 'MAT_076/MAT_077', 'E=100MPa, 초탄성 고려'],
    ['15', 'Mobile.Epoxy', '에폭시+고무 혼합',  'MAT_076',         '점탄성'],
    ['12', 'TAPE_1620MPa', '접착 테이프',       'MAT_076',         '점탄성'],
    ['21', 'SUZY_1520M',   '실리콘/접착제',     'MAT_076/MAT_077', '점탄성/초탄성'],
]
add_table_with_style(
    ['MID', '재료명', '추정 재질', '권장 MAT', '비고'],
    ve_rows
)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('특히 Epoxy_100M (E=100MPa)').bold = True
p.add_run('은 순수 에폭시(E=2000~4000MPa) 대비 매우 낮아, 고무 성분이 상당량 혼합된 것으로 추정됨. ')
p.add_run('MAT_077 (Ogden) 초탄성 모델 적용 검토 필요.')

doc.add_paragraph('')
add_heading_kr('3.2 소성 모델 필요 - 금속류 (3개 추가)', level=2)

pl_rows = [
    ['3',  'Mobile.STS304', 'STS304 스테인리스강', 'MAT_024', '고변형률속도 소성'],
    ['10', 'AL7003H',       '알루미늄 7003-H',     'MAT_024', 'Johnson-Cook 권장'],
    ['20', 'Cu',            '구리',                'MAT_024', '소성 커브 필요'],
]
add_table_with_style(
    ['MID', '재료명', '추정 재질', '권장 MAT', '비고'],
    pl_rows
)

doc.add_paragraph('')
add_heading_kr('3.3 직교이방성 모델 필요 - 섬유 강화 복합재 (7개)', level=2)

ortho_rows = [
    ['2',  'PC-GF30%',      'PC+GF 30%',       'MAT_002/MAT_040', '직교이방성'],
    ['16', 'PC-GF50%',      'PC+GF 50%',       'MAT_002',         '이방성 뚜렷'],
    ['17', 'PC-GF10%',      'PC+GF 10%',       'MAT_024/MAT_076', '등방 근사 가능'],
    ['8',  'PCB_30G',       'PCB (FR-4)',       'MAT_002',         '직교이방성 적층판'],
    ['13', 'PC_QF1035',     'PC',               'MAT_024',         '소성/점탄성'],
    ['18', 'PC_BATT',       '배터리 케이스',     'MAT_024',         '소성'],
    ['19', 'PCGF30_LS3302', 'PC+GF30%',        'MAT_002',         '직교이방성'],
]
add_table_with_style(
    ['MID', '재료명', '추정 재질', '권장 MAT', '비고'],
    ortho_rows
)

doc.add_paragraph('')
add_heading_kr('3.4 탄성으로 충분한 재료 (4개)', level=2)

ok_rows = [
    ['4',  'Ground',    '바닥면 (강체)',   'MAT_001 OK', '변형 불필요'],
    ['14', 'GG8(79G)',  '고릴라 글래스 8', 'MAT_001 OK', '취성 - MAT_ADD_EROSION 추가 가능'],
    ['9',  'Panel',     '디스플레이',      'MAT_001 OK', '취성 근사'],
    ['24', 'Panel1187', '디스플레이',      'MAT_001 OK', '취성 근사'],
]
add_table_with_style(
    ['MID', '재료명', '추정 재질', '현재 상태', '비고'],
    ok_rows
)

doc.add_paragraph('')
add_heading_kr('3.5 고도화 비율 종합', level=2)

summary_rows = [
    ['기본 탄성 (MAT_001)',         '21', '95.5%'],
    ['소성 (MAT_024)',              '1',  '4.5%'],
    ['점탄성/초탄성 (MAT_076/077)', '0',  '0%'],
    ['직교이방성 (MAT_002)',        '0',  '0%'],
]
add_table_with_style(['물성 타입', '개수', '비중'], summary_rows)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('결론: ').bold = True
p.add_run('22개 물성 중 고도화 필요 약 18개(82%), 현재 고도화 완료 1개(4.5%). ')
p.add_run('에폭시 계열은 고무 혼합으로 인해 점탄성/초탄성 모델 적용이 특히 중요함.')

# ============================
# 4. Missing Materials
# ============================
doc.add_page_break()
add_heading_kr('4. 스마트폰 내 누락 재질 분석', level=1)
doc.add_paragraph('실제 스마트폰에 존재하지만 현재 모델에 포함되지 않은 재질 목록.')

add_heading_kr('4.1 구조/섀시 (5개)', level=2)
add_table_with_style(
    ['재질', '사용 위치', '권장 MAT', '중요도'],
    [
        ['티타늄 합금 (Ti-6Al-4V)',  '최신 플래그십 프레임',     'MAT_024',           '상'],
        ['마그네슘 합금 (AZ91D)',    '내부 미드프레임',          'MAT_024',           '중'],
        ['스프링강 (STS301)',        '버튼 스프링/클립',         'MAT_024',           '하'],
        ['사파이어 (Al2O3)',         '카메라 렌즈 커버',         'MAT_001+EROSION',   '중'],
        ['세라믹 (ZrO2)',           '후면 커버 (일부 모델)',     'MAT_001+EROSION',   '하'],
    ]
)

doc.add_paragraph('')
add_heading_kr('4.2 고무/엘라스토머 (7개)', level=2)
doc.add_paragraph('현재 모델에 고무 계열 전무. 충격 흡수에 핵심적. 에폭시+고무 혼합재도 별도 고려 필요.')
add_table_with_style(
    ['재질', '사용 위치', '권장 MAT', '중요도', '비고'],
    [
        ['실리콘 러버',       '방수 가스켓/실링',   'MAT_077 (Ogden)',  '상', 'E~1-10 MPa, PR~0.49'],
        ['TPU (열가소성 PU)', '범퍼/보호 케이스',   'MAT_077/MAT_181', '상', '초탄성+점탄성'],
        ['EPDM 러버',        '버튼 개스킷',        'MAT_077',         '중', 'E~5-20 MPa'],
        ['실리콘 폼 (발포)',  '충격 흡수 패드',     'MAT_057',         '상', '매우 낮은 강성'],
        ['EVA/PE 폼',        '내부 완충재',        'MAT_057/MAT_083', '중', '발포체'],
        ['도전성 고무',       'EMI 가스켓',         'MAT_077',         '하', '전기적 차폐용'],
        ['LSR (액상 실리콘)', '마이크/스피커 씰',   'MAT_077',         '중', '정밀 성형'],
    ]
)

doc.add_paragraph('')
add_heading_kr('4.3 배터리 내부 (8개)', level=2)
doc.add_paragraph('PC_BATT가 케이스만 표현. 배터리 셀 내부 구조 누락. 낙하 시 내부 단락/팽창 평가에 중요.')
add_table_with_style(
    ['재질', '사용 위치', '권장 MAT', '중요도'],
    [
        ['알루미늄 파우치',        '배터리 외장',      'MAT_024',               '상'],
        ['양극재 (LiCoO2)',       '양극 활물질 층',    'MAT_024/MAT_063',       '상'],
        ['음극재 (흑연)',          '음극 활물질 층',    'MAT_063 (Crush Foam)',  '상'],
        ['Cu 집전체 (Cu foil)',   '음극 집전체',       'MAT_024',               '중'],
        ['Al 집전체 (Al foil)',   '양극 집전체',       'MAT_024',               '중'],
        ['세퍼레이터 (PE/PP)',    '양극-음극 분리막',   'MAT_024/MAT_076',       '상'],
        ['전해액',                '리튬이온 전달',      'MAT_NULL+EOS',          '중'],
        ['젤리롤 등가체',         '배터리 셀 전체',     'MAT_126 (Honeycomb)',   '상'],
    ]
)

doc.add_paragraph('')
add_heading_kr('4.4 디스플레이 어셈블리 (6개)', level=2)
add_table_with_style(
    ['재질', '사용 위치', '권장 MAT', '중요도'],
    [
        ['OCA (광학 접착제)',     '디스플레이 접합',    'MAT_076 (점탄성)',  '상'],
        ['편광 필름 (PVA/TAC)',  '편광판',            'MAT_024/MAT_002',  '중'],
        ['OLED 유기층',          '발광 소자',          'MAT_001',          '하'],
        ['터치 센서 (ITO/Ag)',   '터치 패널',          'MAT_001',          '하'],
        ['백플레이트 (STS)',     '디스플레이 보강판',   'MAT_024',          '중'],
        ['디스플레이 폼 쿠션',   '디스플레이 완충',     'MAT_057',          '상'],
    ]
)

doc.add_paragraph('')
add_heading_kr('4.5 열관리 부품 (4개)', level=2)
add_table_with_style(
    ['재질', '사용 위치', '권장 MAT', '중요도'],
    [
        ['그라파이트 시트',   '열 확산',        'MAT_002 (이방성)',  '중'],
        ['열전도 패드 (TIM)', 'IC-히트싱크',    'MAT_076',          '중'],
        ['베이퍼 챔버 (Cu)',  'AP 방열',        'MAT_024',          '하'],
        ['열전도 겔',         '발열 부품 접합',  'MAT_076',          '하'],
    ]
)

doc.add_paragraph('')
add_heading_kr('4.6 카메라/센서 모듈 (5개)', level=2)
add_table_with_style(
    ['재질', '사용 위치', '권장 MAT', '중요도'],
    [
        ['카메라 렌즈 (광학유리)',   '카메라 모듈',         'MAT_001',   '중'],
        ['카메라 하우징 (LCP)',     '렌즈 배럴',           'MAT_024',   '중'],
        ['OIS 스프링 (BeCu)',      '광학 손떨림 보정',     'MAT_024',   '하'],
        ['액츄에이터 자석 (NdFeB)', 'AF/OIS 구동',        'MAT_RIGID', '하'],
        ['IR 필터 (유리)',          '적외선 차단',         'MAT_001',   '하'],
    ]
)

doc.add_paragraph('')
add_heading_kr('4.7 커넥터/소형 기구부품 (7개)', level=2)
add_table_with_style(
    ['재질', '사용 위치', '권장 MAT', '중요도'],
    [
        ['USB-C 커넥터 (STS/인청동)', 'USB 포트',        'MAT_024',   '중'],
        ['SIM 트레이 (STS/Ti)',      'SIM 카드 슬롯',    'MAT_024',   '하'],
        ['스피커 메쉬 (STS)',         '스피커 그릴',       'MAT_001',   '하'],
        ['NFC 안테나 (Cu coil)',     'NFC 통신',         'MAT_001',   '하'],
        ['FPC (폴리이미드+Cu)',      '플렉서블 케이블',    'MAT_002',   '중'],
        ['RF 안테나 (LDS)',          '통신 안테나',        'MAT_001',   '하'],
        ['나사 (STS)',               '체결 부품',          'MAT_024',   '하'],
    ]
)

doc.add_paragraph('')
add_heading_kr('4.8 접착제/필름 (6개)', level=2)
add_table_with_style(
    ['재질', '사용 위치', '권장 MAT', '중요도'],
    [
        ['PSA (감압접착제)',       '후면 커버 접합',      'MAT_076',  '상'],
        ['UV 경화 접착제',        '카메라 렌즈 고정',     'MAT_001',  '하'],
        ['ACF (이방성 도전필름)',  '디스플레이 FPC 접합', 'MAT_076',  '중'],
        ['PET 보호필름',          '내부 절연/보호',       'MAT_024',  '하'],
        ['PI 필름 (Kapton)',     '전기 절연',            'MAT_024',  '중'],
        ['양면 테이프 (VHB)',     '부품 고정',            'MAT_076',  '상'],
    ]
)

# ============================
# 5. Summary
# ============================
doc.add_page_break()
add_heading_kr('5. 종합 요약', level=1)

add_heading_kr('5.1 현재 모델 vs 실제 스마트폰', level=2)
add_table_with_style(
    ['항목', '개수', '비고'],
    [
        ['현재 모델 물성',          '22개',  '-'],
        ['고도화 필요 (현재 물성)',   '18개',  '82%'],
        ['고도화 완료 (현재)',       '1개',   '4.5%'],
        ['---', '---', '---'],
        ['누락: 구조/섀시',         '5개',   '티타늄, 마그네슘 등'],
        ['누락: 고무/폼',           '7개',   '가스켓, 완충재 등'],
        ['누락: 배터리 내부',        '8개',   '셀 내부 구조'],
        ['누락: 디스플레이',         '6개',   'OCA, 폼 쿠션 등'],
        ['누락: 열관리',            '4개',   '그라파이트, TIM 등'],
        ['누락: 카메라/센서',        '5개',   '렌즈, 하우징 등'],
        ['누락: 커넥터/기구',        '7개',   'USB-C, FPC 등'],
        ['누락: 접착/필름',          '6개',   'PSA, VHB 등'],
        ['---', '---', '---'],
        ['누락 재질 합계',          '48개',  '-'],
        ['전체 필요 물성 (추정)',    '~70개', '-'],
    ]
)

doc.add_paragraph('')
add_heading_kr('5.2 우선순위 권장사항', level=2)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('[1순위] 낙하 해석 정확도에 직접 영향').bold = True
doc.add_paragraph('  - 금속 소성 모델: STS304, AL7003H, Cu -> MAT_024 + Johnson-Cook')
doc.add_paragraph('  - 고무/폼 추가: 가스켓, 완충재 -> MAT_077, MAT_057')
doc.add_paragraph('  - 솔더 크리프/소성: BGA, SOLDER -> MAT_024 + 변형률속도')
doc.add_paragraph('  - 에폭시+고무 혼합재 점탄성: Epoxy_100M 등 -> MAT_076/MAT_077')
doc.add_paragraph('  - 배터리 셀 등가 모델: 젤리롤 -> MAT_126 Honeycomb')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('[2순위] 정밀도 향상').bold = True
doc.add_paragraph('  - 에폭시/접착제 점탄성: MAT_076 Prony series')
doc.add_paragraph('  - PCB/GF복합재 직교이방성: MAT_002')
doc.add_paragraph('  - OCA/PSA 접착제: 디스플레이 분리 예측')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('[3순위] 상세 해석').bold = True
doc.add_paragraph('  - 카메라 모듈 상세')
doc.add_paragraph('  - 커넥터/FPC 상세')
doc.add_paragraph('  - 열관리 부품')

# Save
output_path = r'd:\MXDigitalTwinModeller\LSDYNA_Material_Analysis.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
