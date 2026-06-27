"""
MX Digital Twin Modeller - Mechanical ACT Extension 기능 레포트 생성기
Word (.docx) 문서 생성 (한국어 폰트 지원)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ========================================================================
# 한국어 폰트 설정 헬퍼 (SpaceClaim 레포트와 동일)
# ========================================================================

def set_run_font(run, font_name='맑은 고딕', size=10, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_heading_kr(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_run_font(run, '맑은 고딕', size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True)
    return heading


def add_paragraph_kr(doc, text, bold=False, size=10, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, '맑은 고딕', size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet_kr(doc, text, level=0, bold_prefix='', size=10):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, '맑은 고딕', size=size, bold=True)
        run = p.add_run(text)
        set_run_font(run, '맑은 고딕', size=size)
    else:
        run = p.add_run(text)
        set_run_font(run, '맑은 고딕', size=size)
    return p


def add_table_kr(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, '맑은 고딕', size=9, bold=True, color=(255, 255, 255))
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '2E5090'
        })
        shading.append(shading_elm)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, '맑은 고딕', size=9)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def add_screenshot_placeholder(doc, caption=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ 스크린샷: {caption} ]')
    set_run_font(run, '맑은 고딕', size=10, color=(150, 150, 150))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F0F0F0'
    })
    pPr.append(shd)
    return p


# ========================================================================
# 문서 생성
# ========================================================================

def create_report():
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ====================================================================
    # 표지
    # ====================================================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('MX Digital Twin Simulation')
    set_run_font(run, '맑은 고딕', size=28, bold=True, color=(46, 80, 144))

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('ANSYS Mechanical ACT Extension 기능 레포트')
    set_run_font(run, '맑은 고딕', size=18, color=(80, 80, 80))

    doc.add_paragraph()

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_p.add_run('Version 1.1.0')
    set_run_font(run, '맑은 고딕', size=14, color=(120, 120, 120))

    doc.add_paragraph()

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_p.add_run('MX Composites')
    set_run_font(run, '맑은 고딕', size=14, bold=True, color=(46, 80, 144))

    from datetime import date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(date.today().strftime('%Y년 %m월 %d일'))
    set_run_font(run, '맑은 고딕', size=12, color=(120, 120, 120))

    doc.add_page_break()

    # ====================================================================
    # 목차
    # ====================================================================
    add_heading_kr(doc, '목차', level=1)
    toc_items = [
        '1. 개요',
        '2. 시스템 요구사항',
        '3. 툴바 구성',
        '4. MX Digital Twin Simulation 툴바',
        '    4.1 Face Pair NS (접촉면 쌍 Named Selection)',
        '    4.2 Named Selections (방향별 접촉면 NS)',
        '    4.3 Modal Analysis (모드 해석)',
        '    4.4 Add Scenario (과도 해석 시나리오)',
        '    4.5 Post-Process (후처리 뷰어)',
        '    4.6 Export K-File (LS-DYNA 내보내기)',
        '    4.7 Tied Check (접합 검증)',
        '5. MX Material Twin Simulation 툴바',
        '    5.1 Tensile Test (인장시험 재료 캘리브레이션)',
        '6. 보조 모듈',
        '    6.1 Post-Process Viewer (MXPostViewer)',
        '    6.2 Material Calibrator (MaterialCalibrator)',
    ]
    for item in toc_items:
        add_paragraph_kr(doc, item, size=10, space_after=2)

    doc.add_page_break()

    # ====================================================================
    # 1. 개요
    # ====================================================================
    add_heading_kr(doc, '1. 개요', level=1)
    add_paragraph_kr(doc, (
        'MX Digital Twin Simulation은 ANSYS Mechanical용 ACT Extension으로, '
        '접촉면 감지부터 모드 해석, 과도 해석 시나리오 생성, 후처리, LS-DYNA K-File 내보내기까지의 '
        '해석 워크플로를 자동화하는 통합 도구입니다.'
    ))
    add_paragraph_kr(doc, (
        '본 레포트는 Mechanical ACT Extension에 포함된 모든 기능을 설명합니다. '
        'SpaceClaim Add-In에서 생성한 모델을 Mechanical에서 해석하는 전체 파이프라인을 지원합니다.'
    ))

    add_heading_kr(doc, '주요 특징', level=3)
    features = [
        '접촉면 자동 감지 및 Named Selection 생성 (방향별, 쌍별)',
        'Modal Analysis 자동 생성 (모드 수, 최대 주파수 설정)',
        'Transient 해석 시나리오 자동 생성 (CSV 하중 데이터 + Mode Superposition)',
        '후처리: 바디별 변형/응력 평가, Top-N 랭킹, FFT/FRF 뷰어 연동',
        'LS-DYNA K-File 내보내기 (노드, 요소, 파트, 재료, 접촉 포함)',
        '접합 검증: Modal 해석으로 분리된 바디 자동 감지 및 접촉/억제',
        'Material Twin: 인장시험 데이터로 재료 물성 자동 캘리브레이션',
        'Python 기반 WPF UI (ANSYS Mechanical 내장 환경)',
    ]
    for f in features:
        add_bullet_kr(doc, f)

    doc.add_page_break()

    # ====================================================================
    # 2. 시스템 요구사항
    # ====================================================================
    add_heading_kr(doc, '2. 시스템 요구사항', level=1)
    add_table_kr(doc,
        ['항목', '요구사항'],
        [
            ['운영체제', 'Windows 10/11 (64-bit)'],
            ['ANSYS 버전', 'ANSYS 2025 R2 (v252) 이상'],
            ['Mechanical', 'ANSYS Mechanical v252'],
            ['ACT Extension', 'MXSimulator (자동 배포)'],
            ['설치 경로', '%APPDATA%\\Ansys\\v252\\ACT\\extensions\\MXSimulator'],
            ['Python 환경', 'Mechanical 내장 IronPython + WPF'],
            ['후처리 뷰어', 'MXPostViewer.exe (PyInstaller) 또는 Python 3.9+'],
            ['캘리브레이터', 'MaterialCalibrator.exe (PyInstaller) 또는 Python 3.9+'],
        ],
        col_widths=[4, 12]
    )

    doc.add_page_break()

    # ====================================================================
    # 3. 툴바 구성
    # ====================================================================
    add_heading_kr(doc, '3. 툴바 구성', level=1)
    add_paragraph_kr(doc, (
        'MX Digital Twin Simulation은 Mechanical에 2개의 툴바를 추가합니다.'
    ))

    add_table_kr(doc,
        ['툴바', '버튼 수', '기능 요약'],
        [
            ['MX Digital Twin Simulation', '7', 'Face Pair NS, Named Selections, Modal Analysis, Add Scenario, Post-Process, Export K-File, Tied Check'],
            ['MX Material Twin Simulation', '1', 'Tensile Test 재료 캘리브레이션'],
        ],
        col_widths=[5, 2, 9]
    )

    add_screenshot_placeholder(doc, 'Mechanical 툴바 전체 화면')

    doc.add_page_break()

    # ====================================================================
    # 4. MX Digital Twin Simulation 툴바
    # ====================================================================
    add_heading_kr(doc, '4. MX Digital Twin Simulation 툴바', level=1)
    add_paragraph_kr(doc, (
        '해석 모델 준비부터 후처리까지의 전체 워크플로를 지원하는 7개의 기능으로 구성됩니다.'
    ))

    # ─── 4.1 Face Pair NS ───
    add_heading_kr(doc, '4.1 Face Pair NS (접촉면 쌍 Named Selection)', level=2)
    add_paragraph_kr(doc, (
        '바디 간 접촉면을 자동으로 감지하여 A-side/B-side 쌍으로 Named Selection을 생성합니다. '
        'Include/Exclude 키워드 필터링으로 대상 바디를 선택할 수 있습니다.'
    ))

    add_heading_kr(doc, '워크플로', level=3)
    steps = [
        'Include 키워드 입력 (대상 바디 필터, "all" = 전체)',
        'Exclude 키워드 입력 (제외할 바디)',
        'Tolerance (mm) 설정',
        '"Find Contact Face Pairs" 실행',
        '감지된 쌍 목록에서 체크/언체크로 선택',
        'NS 생성 모드 선택 후 "Create A/B Named Selections" 실행',
    ]
    for i, s in enumerate(steps, 1):
        add_bullet_kr(doc, s, bold_prefix=f'{i}. ')

    add_heading_kr(doc, 'NS 생성 모드', level=3)
    add_table_kr(doc,
        ['모드', '생성 결과', '설명'],
        [
            ['Per Pair', 'FaceA_001/FaceB_001 (쌍마다)', '각 접촉 쌍에 대해 개별 NS 생성'],
            ['Merged', 'Contact_A/Contact_B (전체 합산)', '모든 A-면을 하나의 NS, 모든 B-면을 하나의 NS로 병합'],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading_kr(doc, '감지 알고리즘', level=3)
    feat_list = [
        '법선 벡터 반평행 검사: dot(nA, nB) < -0.8 (143도 이상)',
        '거리 검사: 면 중심점 간 법선 방향 투영 거리 < Tolerance',
        '중복 제거: 기하학적 키 기반 (A,B)/(B,A) 중복 방지',
        '동일 바디 쌍 자동 제외',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('Include: ', '대상 바디 키워드 (쉼표 구분, "all" = 전체 바디)'),
        ('Exclude: ', '제외 바디 키워드 (바디 이름에 포함 시 제외)'),
        ('Tolerance: ', '접촉 판정 허용치 (mm, 기본값 0.1)'),
        ('Merge Prefix: ', 'Merged 모드 시 NS 접두사 (기본값 "Contact")'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_screenshot_placeholder(doc, 'Face Pair NS 대화창 및 감지 결과')

    doc.add_page_break()

    # ─── 4.2 Named Selections ───
    add_heading_kr(doc, '4.2 Named Selections (방향별 접촉면 NS)', level=2)
    add_paragraph_kr(doc, (
        '접촉면을 법선 방향별(+X, -X, +Y, -Y, +Z, -Z)로 분류하여 '
        'Cap_Upper, Cap_Lower 등의 의미론적 이름으로 Named Selection을 생성합니다. '
        '캡 진동 해석의 하중 적용 면을 자동으로 식별합니다.'
    ))

    add_heading_kr(doc, '방향 분류', level=3)
    add_table_kr(doc,
        ['법선 방향', 'NS 이름', '의미'],
        [
            ['+Z', 'Cap_Upper_001', '상면 (캡 상단)'],
            ['-Z', 'Cap_Lower_001', '하면 (캡 하단)'],
            ['+Y', 'Cap_Front_001', '전면'],
            ['-Y', 'Cap_Back_001', '후면'],
            ['+X', 'Cap_Right_001', '우측면'],
            ['-X', 'Cap_Left_001', '좌측면'],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('Keywords: ', '대상 바디 키워드 (빈칸 = 전체 바디)'),
        ('Tolerance: ', '접촉 판정 허용치 (mm, 기본값 0.1)'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        '체크박스로 개별 면 선택/제외 가능',
        '방향별 자동 그룹화 및 넘버링 (001, 002, ...)',
        '기존 NS와 번호 충돌 방지 (자동 증가)',
        '상세 로그 표시 (바디별 감지 결과)',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Named Selections 대화창')

    doc.add_page_break()

    # ─── 4.3 Modal Analysis ───
    add_heading_kr(doc, '4.3 Modal Analysis (모드 해석)', level=2)
    add_paragraph_kr(doc, (
        'Modal Analysis를 자동으로 생성합니다. 모드 수와 최대 주파수를 설정하면 '
        'Mechanical에 Modal 해석을 추가하고 해석 설정을 자동으로 구성합니다.'
    ))

    add_heading_kr(doc, '주요 파라미터', level=3)
    add_table_kr(doc,
        ['파라미터', '기본값', '설명'],
        [
            ['Number of Modes', '20', '찾을 최대 모드 수'],
            ['Max Frequency', '1000 Hz', '주파수 검색 범위 상한'],
        ],
        col_widths=[4, 3, 9]
    )

    add_heading_kr(doc, '동작', level=3)
    feat_list = [
        'model.AddModalAnalysis() 로 Modal 해석 자동 생성',
        'MaximumModesToFind, ModalRangeMaximum 자동 설정',
        'LimitSearchToRange = True (지정 범위 내 검색)',
        '기존 Modal 해석이 있으면 이름 표시 (중복 방지)',
        '생성 후 Fixed Support 추가 안내',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Modal Analysis 대화창')

    # ─── 4.4 Add Scenario ───
    add_heading_kr(doc, '4.4 Add Scenario (과도 해석 시나리오)', level=2)
    add_paragraph_kr(doc, (
        'Cap_* Named Selection에 시간-하중 CSV 데이터를 적용하여 '
        'Transient Structural 해석을 자동으로 생성합니다. '
        'Mode Superposition 또는 Full Transient 방식을 선택할 수 있습니다.'
    ))

    add_heading_kr(doc, '워크플로', level=3)
    steps = [
        'Cap_* NS 목록 자동 로드 (Named Selections 기능에서 생성된 NS)',
        '각 NS에 대해 하중 방향 선택 (Positive + / Negative -)',
        'CSV 파일 선택 (Time(s), Force(N) 2컬럼)',
        'End Time, Time Step 설정',
        '해석 방식 선택 (Mode Superposition / Full Transient)',
        '"Create Scenario" 실행',
    ]
    for i, s in enumerate(steps, 1):
        add_bullet_kr(doc, s, bold_prefix=f'{i}. ')

    add_heading_kr(doc, '주요 파라미터', level=3)
    add_table_kr(doc,
        ['파라미터', '기본값', '설명'],
        [
            ['End Time', '0.02 s', '해석 종료 시간'],
            ['Time Step', '0.0001 s', '시간 간격 (min = max = step)'],
            ['Method', 'Mode Superposition', '해석 방식 (Modal 자동 연결)'],
        ],
        col_widths=[4, 3, 9]
    )

    add_heading_kr(doc, '특징', level=3)
    feat_list = [
        'NS별 하중 방향 개별 설정: NS 이름으로 기본 방향 자동 감지 (Lower→-Z, Upper→+Z 등)',
        '사용자 Override: +/- RadioButton으로 방향 수동 전환',
        'Mode Superposition: 기존 Modal 해석 자동 연결 (model.Link)',
        '자동 네이밍: Transient_Scenario_1, _2, ... 순차 번호',
        'CSV 파싱: 헤더 자동 스킵, 숫자 데이터만 추출',
        '축별 하중 적용: NS 이름에서 축(X/Y/Z) 자동 판별하여 ForceComponent 설정',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Add Scenario 대화창')

    doc.add_page_break()

    # ─── 4.5 Post-Process ───
    add_heading_kr(doc, '4.5 Post-Process (후처리 뷰어)', level=2)
    add_paragraph_kr(doc, (
        'Transient 해석 결과에서 바디별 Total Deformation과 Equivalent Stress를 평가하고, '
        'Top-N 랭킹을 생성한 후, 시간이력 CSV로 내보내어 외부 PyQt5 뷰어를 실행합니다.'
    ))

    add_heading_kr(doc, '2단계 워크플로', level=3)
    add_table_kr(doc,
        ['단계', '기능', '설명'],
        [
            ['① Add Results & Evaluate', '결과 추가 + 평가', '바디별 TotalDeformation + EquivalentStress 추가 → 전체 평가 → Top-N 랭킹'],
            ['② Export CSV + Launch Viewer', 'CSV 내보내기 + 뷰어', '시간이력 CSV 내보내기 → metadata.json 생성 → MXPostViewer 실행'],
        ],
        col_widths=[4, 4, 8]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('Analysis: ', 'Transient 해석 선택 (드롭다운)'),
        ('Top-N: ', '랭킹할 바디 수 (기본 10)'),
        ('Op. Freq: ', '운영 주파수 Hz (뷰어에 전달)'),
        ('Thresh Red: ', '적색 경고 임계값 (mm)'),
        ('Thresh Yellow: ', '황색 경고 임계값 (mm)'),
        ('Output Folder: ', 'CSV 및 metadata.json 저장 경로'),
        ('Force CSV: ', '입력 하중 CSV (FRF 계산용, 선택사항)'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '① Add Results & Evaluate 상세', level=3)
    feat_list = [
        '전체 바디 자동 감지 (model.Geometry.GetChildren)',
        '바디별 TotalDeformation + EquivalentStress 결과 자동 추가',
        'sol.EvaluateAllResults() 로 일괄 평가',
        'MaximumOfMaximumOverTime 기준 랭킹',
        '해석 미풀이 상태 감지 및 안내 메시지',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_heading_kr(doc, '② Export CSV + Launch Viewer 상세', level=3)
    feat_list = [
        '각 랭킹 바디의 시간이력 데이터를 TXT 파일로 내보내기',
        'metadata.json 생성 (해석명, 주파수, 임계값, 바디 목록)',
        'MXPostViewer.exe 자동 실행 (PyInstaller 번들)',
        'Python fallback: venv → PATH → 공통 경로 순서로 탐색',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Post-Process 대화창')

    doc.add_page_break()

    # ─── 4.6 Export K-File ───
    add_heading_kr(doc, '4.6 Export K-File (LS-DYNA 내보내기)', level=2)
    add_paragraph_kr(doc, (
        'Mechanical 모델의 메쉬를 LS-DYNA .k 파일 형식으로 내보냅니다. '
        '노드, 요소, 파트, 재료, Named Selection, Contact Region을 모두 포함합니다.'
    ))

    add_heading_kr(doc, '내보내기 항목', level=3)
    add_table_kr(doc,
        ['LS-DYNA 키워드', '내용', '설명'],
        [
            ['*NODE', '노드 좌표', 'NID, X, Y, Z (mm 또는 m 단위)'],
            ['*ELEMENT_SOLID/SHELL', '요소 연결성', 'Tet4/10, Hex8/20, Wedge6/15, Shell3/4 지원'],
            ['*PART', '파트 정의', '바디별 자동 PID 할당'],
            ['*SECTION_SOLID/SHELL', '단면 정의', '요소 유형별 ELFORM 자동 설정'],
            ['*MAT_ELASTIC', '재료 물성', 'E, ν, ρ (미할당 시 기본 강재)'],
            ['*SET_NODE_TITLE', 'Named Selection', '기하학적 면 매칭으로 노드 집합 생성'],
            ['*CONTACT_*', '접촉 조건', 'Master/Slave 자동 할당 (메쉬 크기 기반)'],
            ['*SET_SEGMENT_TITLE', '접촉 면 집합', '접촉면 요소 면 매칭'],
        ],
        col_widths=[4, 3, 9]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('Output Path: ', '.k 파일 저장 경로 (SaveFileDialog)'),
        ('Unit System: ', 'mm-tonne-s (LS-DYNA 표준) 또는 SI (m-kg-s)'),
        ('Named Selections: ', '*SET_NODE_TITLE로 내보내기 여부'),
        ('Materials: ', '*MAT_ELASTIC로 내보내기 여부'),
        ('Contact Regions: ', '*CONTACT_* + *SET_SEGMENT_TITLE로 내보내기 여부'),
        ('Geo Tolerance: ', '기하학적 면 매칭 허용치 (mm, 기본 0.1)'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, '요소 유형 변환', level=3)
    add_table_kr(doc,
        ['Mechanical 요소', '노드 수', 'LS-DYNA 변환', 'ELFORM'],
        [
            ['Tet4', '4', '*ELEMENT_SOLID (degenerate hex)', '10'],
            ['Tet10', '10', 'Tet4로 저감 (corner nodes)', '10'],
            ['Hex8', '8', '*ELEMENT_SOLID', '1'],
            ['Hex20', '20', 'Hex8로 저감', '1'],
            ['Wedge6', '6', '*ELEMENT_SOLID (degenerate)', '15'],
            ['Shell3', '3', '*ELEMENT_SHELL', '-'],
            ['Shell4', '4', '*ELEMENT_SHELL', '-'],
        ],
        col_widths=[4, 2, 5, 2]
    )

    add_screenshot_placeholder(doc, 'Export K-File 대화창')

    doc.add_page_break()

    # ─── 4.7 Tied Check ───
    add_heading_kr(doc, '4.7 Tied Check (접합 검증)', level=2)
    add_paragraph_kr(doc, (
        'Modal 해석을 통해 분리된(비접합) 바디를 자동으로 감지합니다. '
        '강체 모드(Rigid Body Mode, RBM)를 분석하여 접촉이 누락된 바디를 식별하고, '
        '자동 접촉 생성 또는 바디 억제를 수행합니다.'
    ))

    add_heading_kr(doc, '4단계 워크플로', level=3)
    add_table_kr(doc,
        ['단계', '버튼', '기능'],
        [
            ['1', '① Create & Solve Modal', 'Modal 해석 자동 생성 및 풀이 → RBM 감지'],
            ['2', '② Create Contacts', '분리된 바디에 Face-to-Face 접촉 자동 생성'],
            ['3', '③ Suppress Checked', '체크된 분리 바디 억제 (해석 제외)'],
            ['4', '④ Re-Run Modal', '수정 후 Modal 재풀이 → 재검증 (반복)'],
        ],
        col_widths=[1, 5, 10]
    )

    add_heading_kr(doc, '주요 파라미터', level=3)
    params = [
        ('Max Modes: ', '검색할 최대 모드 수 (기본 20)'),
        ('RBM Threshold: ', '강체 모드 판정 주파수 임계값 (Hz, 기본 0.1)'),
        ('Tolerance: ', '접촉 생성 허용치 (mm, 기본 0.01)'),
    ]
    for prefix, desc in params:
        add_bullet_kr(doc, desc, bold_prefix=prefix)

    add_heading_kr(doc, 'RBM 감지 알고리즘', level=3)
    feat_list = [
        'Modal 해석 풀이 후 전체 모드 주파수 읽기',
        'Threshold (Hz) 미만 모드를 RBM으로 분류',
        'Contact Region 연결성 분석 (BFS 기반 그래프 탐색)',
        '주 연결 컴포넌트가 아닌 바디를 분리 후보로 식별',
        '분리 후보 바디에 대해 변형 에너지 비율 검사',
        '6개 이하 RBM → 전역 RBM (정상), 초과 시 개별 바디 분석',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_heading_kr(doc, '반복 검증', level=3)
    feat_list = [
        'Restore All: 억제된 바디 모두 복원',
        'Iteration 카운터: 반복 횟수 및 누적 억제 수 표시',
        '① → ②/③ → ④ 순환으로 모든 바디 접합 확인',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_screenshot_placeholder(doc, 'Tied Check 대화창 (분리 바디 감지 결과)')

    doc.add_page_break()

    # ====================================================================
    # 5. MX Material Twin Simulation 툴바
    # ====================================================================
    add_heading_kr(doc, '5. MX Material Twin Simulation 툴바', level=1)
    add_paragraph_kr(doc, (
        'Material Twin 기능은 실험 데이터로부터 재료 물성을 자동으로 역추정(캘리브레이션)합니다.'
    ))

    # ─── 5.1 Tensile Test ───
    add_heading_kr(doc, '5.1 Tensile Test (인장시험 재료 캘리브레이션)', level=2)
    add_paragraph_kr(doc, (
        '인장시험의 힘-변위 데이터로부터 재료 물성(영률, 포아송비 등)을 '
        '자동으로 캘리브레이션합니다. 5단계 마법사 형태로 구성됩니다.'
    ))

    add_heading_kr(doc, '5단계 워크플로', level=3)
    add_table_kr(doc,
        ['단계', '기능', '설명'],
        [
            ['Step 1', 'Specimen Detection', '시편 기하 정보 감지 (Workbench Parameter / YAML / JSON / Manual)'],
            ['Step 2', 'CSV Data Input', '인장시험 CSV 로드 (displacement_mm, force_N)'],
            ['Step 3', 'Material Properties', '포아송비(ν), 밀도(ρ) 입력 (또는 기본값 사용)'],
            ['Step 4', 'Elastic Calibration', "Young's Modulus (E) 자동 캘리브레이션"],
            ['Step 5', 'Engineering Data', 'Mechanical Engineering Data에 재료 자동 생성'],
        ],
        col_widths=[2, 4, 10]
    )

    add_heading_kr(doc, 'Step 1: 시편 감지 우선순위', level=3)
    add_table_kr(doc,
        ['우선순위', '소스', '설명'],
        [
            ['1', 'Workbench Parameters', 'P1_GaugeLength, P2_GaugeWidth, P3_Thickness'],
            ['2', 'specimen.yaml', 'SpaceClaim에서 시편 생성 시 자동 저장된 파일'],
            ['3', '.specimen.json', 'STEP 내보내기 시 생성된 JSON 메타데이터'],
            ['4', 'Manual Input', '사용자 직접 입력 (대화창)'],
        ],
        col_widths=[2, 4, 10]
    )

    add_heading_kr(doc, '캘리브레이션 프로세스', level=3)
    feat_list = [
        'CSV에서 힘-변위 곡선 로드',
        '시편 단면적 계산 (GaugeWidth x Thickness)',
        '공칭 응력(σ) = Force / Area, 공칭 변형률(ε) = Displacement / GaugeLength',
        '선형 영역 자동 감지 (응력-변형률 곡선의 기울기 안정 구간)',
        "Young's Modulus (E) = Δσ / Δε (선형 영역)",
        'Engineering Data에 Isotropic Elasticity 재료 자동 생성',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_heading_kr(doc, '지원 재료 모델 (확장 예정)', level=3)
    add_table_kr(doc,
        ['재료 모델', '입력 데이터', '상태'],
        [
            ['Isotropic Elasticity (E, ν)', '인장시험 CSV', '구현 완료'],
            ['Plastic Hardening', '응력-변형률 곡선', '확장 예정'],
            ['Viscoelasticity (Prony Series)', 'DMA 데이터', '확장 예정'],
            ['Hyperelastic + Damping', '고무/테이프 시험 데이터', '확장 예정'],
        ],
        col_widths=[5, 5, 3]
    )

    add_screenshot_placeholder(doc, 'Material Twin 대화창 (5단계 마법사)')

    doc.add_page_break()

    # ====================================================================
    # 6. 보조 모듈
    # ====================================================================
    add_heading_kr(doc, '6. 보조 모듈', level=1)
    add_paragraph_kr(doc, (
        'Mechanical ACT Extension과 연동되는 외부 실행 모듈입니다. '
        'PyInstaller로 빌드된 독립 실행 파일(.exe) 또는 Python 스크립트로 실행됩니다.'
    ))

    # ─── 6.1 MXPostViewer ───
    add_heading_kr(doc, '6.1 Post-Process Viewer (MXPostViewer)', level=2)
    add_paragraph_kr(doc, (
        'PyQt5 기반 독립 후처리 뷰어입니다. Mechanical의 Post-Process 기능에서 '
        '내보낸 시간이력 CSV와 metadata.json을 읽어 시각화합니다.'
    ))

    add_heading_kr(doc, '주요 기능', level=3)
    feat_list = [
        'Time History 그래프: 바디별 변형/응력 시간이력 표시',
        'FFT 분석: 시간 영역 → 주파수 영역 변환 (numpy.fft)',
        'FRF 계산: 입력 하중 대비 응답 비율 (Force CSV 필요)',
        'Top-N 바디 랭킹: 최대 변형 기준 순위 표시',
        '신호등 시스템: Red/Yellow/Green 경고 표시 (임계값 기반)',
        'metadata.json: 해석 설정 및 바디 정보 자동 로드',
    ]
    for f in feat_list:
        add_bullet_kr(doc, f)

    add_heading_kr(doc, '실행 방법', level=3)
    add_table_kr(doc,
        ['방법', '설명'],
        [
            ['MXPostViewer.exe', 'PyInstaller 번들 (별도 Python 불필요)'],
            ['python runner.py', 'Python + PyQt5 환경에서 직접 실행'],
            ['Mechanical 연동', 'Post-Process ② 버튼에서 자동 실행'],
        ],
        col_widths=[5, 11]
    )

    add_screenshot_placeholder(doc, 'MXPostViewer 후처리 화면')

    # ─── 6.2 MaterialCalibrator ───
    add_heading_kr(doc, '6.2 Material Calibrator (MaterialCalibrator)', level=2)
    add_paragraph_kr(doc, (
        'PyInstaller 기반 독립 재료 캘리브레이션 도구입니다. '
        'Mechanical 외부에서도 인장시험 데이터로 재료 물성을 추출할 수 있습니다.'
    ))

    add_heading_kr(doc, '구성 파일', level=3)
    add_table_kr(doc,
        ['파일', '역할'],
        [
            ['runner.py', '메인 실행 스크립트'],
            ['elastic_calibrator.py', '탄성 캘리브레이션 로직'],
            ['utils/yaml_parser.py', 'YAML 파서 (시편 메타데이터)'],
            ['utils/csv_parser.py', 'CSV 파서 (시험 데이터)'],
        ],
        col_widths=[5, 11]
    )

    add_screenshot_placeholder(doc, 'MaterialCalibrator 실행 화면')

    doc.add_page_break()

    # ====================================================================
    # 기능 요약 테이블
    # ====================================================================
    add_heading_kr(doc, '기능 요약', level=1)
    add_table_kr(doc,
        ['No.', '툴바', '기능명', '설명'],
        [
            ['1', 'Simulation', 'Face Pair NS', '접촉면 쌍 감지 → A/B NS 생성'],
            ['2', 'Simulation', 'Named Selections', '방향별 접촉면 NS 생성 (Cap_*)'],
            ['3', 'Simulation', 'Modal Analysis', '모드 해석 자동 생성'],
            ['4', 'Simulation', 'Add Scenario', 'CSV + NS → Transient 해석 생성'],
            ['5', 'Simulation', 'Post-Process', '바디별 평가 + Top-N + 뷰어 연동'],
            ['6', 'Simulation', 'Export K-File', 'LS-DYNA .k 파일 내보내기'],
            ['7', 'Simulation', 'Tied Check', 'Modal RBM으로 분리 바디 감지/수정'],
            ['8', 'Material Twin', 'Tensile Test', '인장시험 데이터 → 재료 물성 캘리브레이션'],
        ],
        col_widths=[1, 3, 4, 8]
    )

    doc.add_paragraph()

    add_heading_kr(doc, '보조 모듈', level=3)
    add_table_kr(doc,
        ['No.', '모듈', '설명'],
        [
            ['A', 'MXPostViewer', 'PyQt5 후처리 뷰어 (FFT, FRF, Top-N)'],
            ['B', 'MaterialCalibrator', '독립 재료 캘리브레이션 도구'],
        ],
        col_widths=[1, 5, 10]
    )

    # ====================================================================
    # 저장
    # ====================================================================
    output_path = os.path.join(os.path.dirname(__file__),
                               'MXDigitalTwinModeller_Mechanical_Report.docx')
    doc.save(output_path)
    print(f'Report saved: {output_path}')
    return output_path


if __name__ == '__main__':
    create_report()
